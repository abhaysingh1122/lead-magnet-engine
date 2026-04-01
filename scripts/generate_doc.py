"""
generate_doc.py
Converts a markdown file into a branded Abhay Singh Word document (.docx).

Usage:
    python scripts/generate_doc.py \
        --title "Title" \
        --type "Playbook" \
        --subtitle "Subtitle" \
        --content output/file.md \
        --output output/file.docx \
        --images-dir output
"""

import argparse
import os
import re
import sys
from pathlib import Path
try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    sys.exit("python-docx not installed. Run: pip install python-docx")

# ── Brand colours ──────────────────────────────────────────────────────────────
BLUE        = RGBColor(0x00, 0x56, 0xA7)
BLUE_DARK   = RGBColor(0x0A, 0x19, 0x33)
TEXT        = RGBColor(0x1A, 0x1A, 0x1A)
MUTED       = RGBColor(0x55, 0x55, 0x55)
BORDER_GREY = RGBColor(0xDD, 0xDD, 0xDD)
CODE_TEXT   = RGBColor(0xD4, 0xD4, 0xD4)   # light grey on dark bg

# Callout palette (bg_hex, border_hex) by emoji group
CALLOUT_STYLES = {
    "warning": ("FFF8ED", "F59E0B"),   # ⚠️ 🚨
    "insight": ("F0FDF4", "10B981"),   # 💡 ✅
    "data":    ("EBF4FF", "3B82F6"),   # 📊 📈 📅
}

_CTA_SENTINEL = "want help running"
BOOKING_URL   = "https://calendly.com/abhaysinghnagarkoti11/new-meeting"
FONT_NAME     = "Calibri"   # Word-safe; closest to Bricolage Grotesque available


# ── XML helpers ────────────────────────────────────────────────────────────────


def _shade_cell(cell, hex_colour):
    """Fill a table cell background with a hex colour."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_colour.lstrip("#"))
    tcPr.append(shd)


def _set_cell_border(cell, side, sz, hex_colour):
    """Set one border on a table cell (side = 'left'/'top'/'right'/'bottom')."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tbl_borders = tcPr.find(qn("w:tcBorders"))
    if tbl_borders is None:
        tbl_borders = OxmlElement("w:tcBorders")
        tcPr.append(tbl_borders)
    border = OxmlElement(f"w:{side}")
    border.set(qn("w:val"),   "single")
    border.set(qn("w:sz"),    str(sz))
    border.set(qn("w:space"), "0")
    border.set(qn("w:color"), hex_colour.lstrip("#"))
    tbl_borders.append(border)


def _set_table_borders_none(tbl):
    """Remove all visible borders from a table (used for callout / image boxes)."""
    tblPr = tbl._tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl._tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "none")
        el.set(qn("w:sz"),    "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        tblBorders.append(el)
    tblPr.append(tblBorders)


def _para_spacing(para, before=0, after=0):
    pPr = para._p.get_or_add_pPr()
    # Remove any existing w:spacing to avoid duplicate XML elements
    existing = pPr.find(qn("w:spacing"))
    if existing is not None:
        pPr.remove(existing)
    spc = OxmlElement("w:spacing")
    spc.set(qn("w:before"), str(before))
    spc.set(qn("w:after"),  str(after))
    pPr.append(spc)


def _set_run_font(run, name=FONT_NAME):
    run.font.name = name
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"),    name)
    rFonts.set(qn("w:hAnsi"),    name)
    rFonts.set(qn("w:eastAsia"), name)
    rFonts.set(qn("w:cs"),       name)


def _set_courier_font(run):
    """Set Courier New on a run without touching any other run properties."""
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rFonts.set(qn(attr), "Courier New")


# ── Inline text parser (bold, links) ──────────────────────────────────────────

def _add_inline(para, text, base_size=10.5, colour=None, bold=False, italic=False):
    """
    Parse **bold**, *italic*, [text](url) and plain text,
    then add runs to `para`.
    """
    # Strip em dashes just in case
    text = text.replace("\u2014", " - ").replace("\u2013", "-")

    pattern = re.compile(
        r'(\*\*(.+?)\*\*)'         # **bold**
        r'|(\*(.+?)\*)'            # *italic*
        r'|(\[([^\]]+)\]\(([^)]+)\))'  # [link text](url)
        r'|(`.+?`)'                # `code`
    )

    pos = 0
    for m in pattern.finditer(text):
        # plain text before this match
        if m.start() > pos:
            run = para.add_run(text[pos:m.start()])
            run.font.size = Pt(base_size)
            run.bold = bold
            run.italic = italic
            if colour:
                run.font.color.rgb = colour
            _set_run_font(run)

        if m.group(1):      # **bold**
            run = para.add_run(m.group(2))
            run.font.size = Pt(base_size)
            run.bold = True
            if colour:
                run.font.color.rgb = colour
            _set_run_font(run)

        elif m.group(3):    # *italic*
            run = para.add_run(m.group(4))
            run.font.size = Pt(base_size)
            run.italic = True
            if colour:
                run.font.color.rgb = colour
            _set_run_font(run)

        elif m.group(5):    # [link text](url)
            link_text = m.group(6)
            run = para.add_run(link_text)
            run.font.size = Pt(base_size)
            run.font.color.rgb = BLUE
            run.underline = True
            _set_run_font(run)

        elif m.group(8):    # `code`
            run = para.add_run(m.group(8)[1:-1])
            run.font.size = Pt(base_size - 0.5)
            if colour:
                run.font.color.rgb = colour
            _set_courier_font(run)

        pos = m.end()

    # trailing plain text
    if pos < len(text):
        run = para.add_run(text[pos:])
        run.font.size = Pt(base_size)
        run.bold = bold
        run.italic = italic
        if colour:
            run.font.color.rgb = colour
        _set_run_font(run)


# ── Cover page ────────────────────────────────────────────────────────────────

def _add_cover(doc, title, type_, subtitle):
    # Type label
    p = doc.add_paragraph()
    _para_spacing(p, before=0, after=60)
    run = p.add_run(f"ABHAY SINGH  •  {type_.upper()}")
    run.font.size = Pt(7.5)
    run.font.bold = True
    run.font.color.rgb = MUTED
    _set_run_font(run)

    # Rule
    rule = doc.add_paragraph()
    _para_spacing(rule, before=0, after=80)
    pPr  = rule._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "12")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "0056A7")
    pBdr.append(bot)
    pPr.append(pBdr)

    # Title
    p = doc.add_paragraph()
    _para_spacing(p, before=0, after=100)
    run = p.add_run(title)
    run.font.size  = Pt(26)
    run.font.bold  = True
    run.font.color.rgb = BLUE_DARK
    _set_run_font(run)

    # Subtitle
    p = doc.add_paragraph()
    _para_spacing(p, before=0, after=80)
    run = p.add_run(subtitle)
    run.font.size  = Pt(11)
    run.font.color.rgb = MUTED
    _set_run_font(run)

    # Divider
    rule2 = doc.add_paragraph()
    _para_spacing(rule2, before=0, after=200)
    pPr2  = rule2._p.get_or_add_pPr()
    pBdr2 = OxmlElement("w:pBdr")
    bot2  = OxmlElement("w:bottom")
    bot2.set(qn("w:val"),   "single")
    bot2.set(qn("w:sz"),    "4")
    bot2.set(qn("w:space"), "1")
    bot2.set(qn("w:color"), "DDDDDD")
    pBdr2.append(bot2)
    pPr2.append(pBdr2)


# ── Callout block ─────────────────────────────────────────────────────────────

def _add_callout(doc, emoji, content):
    emoji = emoji.strip()
    if emoji in ("⚠️", "🚨"):
        style_key = "warning"
    elif emoji in ("💡", "✅"):
        style_key = "insight"
    else:
        style_key = "data"

    bg_hex, border_hex = CALLOUT_STYLES[style_key]

    # Single-cell table to simulate tinted box
    tbl = doc.add_table(rows=1, cols=1)
    _set_table_borders_none(tbl)
    cell = tbl.cell(0, 0)
    _shade_cell(cell, bg_hex)
    _set_cell_border(cell, "left", 18, border_hex)  # thick left border

    # Clear default empty para; add our own
    cell.paragraphs[0].clear()
    p = cell.paragraphs[0]
    _para_spacing(p, before=60, after=60)

    run_e = p.add_run(emoji + "  ")
    run_e.font.size = Pt(10)
    _set_run_font(run_e)

    _add_inline(p, content, base_size=10, colour=TEXT)

    # Spacing after table
    spacer = doc.add_paragraph()
    _para_spacing(spacer, before=0, after=60)


# ── Horizontal rule ───────────────────────────────────────────────────────────

def _add_hr(doc):
    p = doc.add_paragraph()
    _para_spacing(p, before=80, after=80)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "4")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "DDDDDD")
    pBdr.append(bot)
    pPr.append(pBdr)


# ── Image block ───────────────────────────────────────────────────────────────

def _add_image_block(doc, filename, images_dir):
    if not images_dir:
        return
    img_path = Path(images_dir) / Path(filename).name
    if not img_path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_spacing(p, before=80, after=80)
    run = p.add_run()
    run.add_picture(str(img_path), width=Inches(5.5))


# ── Table ─────────────────────────────────────────────────────────────────────

def _add_table(doc, rows_data):
    """rows_data: list of lists of strings. First row is the header."""
    if not rows_data:
        return
    ncols = max(len(r) for r in rows_data)
    tbl   = doc.add_table(rows=len(rows_data), cols=ncols)
    tbl.style = "Table Grid"

    for i, row_cells in enumerate(rows_data):
        row = tbl.rows[i]
        for j, cell_text in enumerate(row_cells):
            if j >= ncols:
                break
            cell = row.cells[j]
            cell.paragraphs[0].clear()
            p = cell.paragraphs[0]
            _para_spacing(p, before=40, after=40)
            if i == 0:
                # Header row: blue background, white text
                _shade_cell(cell, "0056A7")
                run = p.add_run(cell_text.strip())
                run.font.size  = Pt(9)
                run.font.bold  = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                _set_run_font(run)
            else:
                _add_inline(p, cell_text.strip(), base_size=9.5)

    spacer = doc.add_paragraph()
    _para_spacing(spacer, before=0, after=60)


# ── CTA section ───────────────────────────────────────────────────────────────

def _add_cta(doc):
    _add_hr(doc)

    p = doc.add_paragraph()
    _para_spacing(p, before=80, after=40)
    run = p.add_run("Want qualified leads without the guesswork?")
    run.font.size  = Pt(13)
    run.font.bold  = True
    run.font.color.rgb = BLUE
    _set_run_font(run)

    p2 = doc.add_paragraph()
    _para_spacing(p2, before=0, after=60)
    run2 = p2.add_run(
        "Free 30-min strategy call. We will audit your lead gen setup and give you "
        "a roadmap you can act on that week."
    )
    run2.font.size  = Pt(10)
    run2.font.color.rgb = MUTED
    _set_run_font(run2)

    p3 = doc.add_paragraph()
    _para_spacing(p3, before=0, after=80)
    run3 = p3.add_run("Book a strategy call: " + BOOKING_URL)
    run3.font.size  = Pt(10)
    run3.font.bold  = True
    run3.font.color.rgb = BLUE
    run3.underline  = True
    _set_run_font(run3)


# ── Code block ────────────────────────────────────────────────────────────────

def _add_code_block(doc, code_lines):
    """Render a fenced code block: dark charcoal bg, Courier New, blue left border."""
    tbl = doc.add_table(rows=1, cols=1)
    _set_table_borders_none(tbl)
    cell = tbl.cell(0, 0)
    _shade_cell(cell, "1E1E2E")
    _set_cell_border(cell, "left", 12, "3B82F6")

    for idx, code_line in enumerate(code_lines):
        p = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
        before = 80 if idx == 0 else 0
        after  = 80 if idx == len(code_lines) - 1 else 0
        _para_spacing(p, before=before, after=after)

        run = p.add_run(code_line if code_line else " ")
        run.font.size = Pt(8.5)
        run.font.color.rgb = CODE_TEXT
        _set_courier_font(run)

    # If code_lines was empty, still add padding to the single default para
    if not code_lines:
        _para_spacing(cell.paragraphs[0], before=80, after=80)

    spacer = doc.add_paragraph()
    _para_spacing(spacer, before=0, after=80)


# ── Markdown → docx ───────────────────────────────────────────────────────────

def _is_table_row(line):
    return line.strip().startswith("|") and line.strip().endswith("|")


def _is_table_separator(line):
    return bool(re.match(r'^\s*\|[\s\-|:]+\|\s*$', line))


def _parse_table_row(line):
    cells = line.strip().strip("|").split("|")
    return [c.strip() for c in cells]


def parse_and_add(doc, text, images_dir=None):
    """Parse markdown text and add content to the docx."""
    lines = text.split("\n")
    i = 0
    list_context = None   # "bullet" | "numbered" | "checkbox" | None

    def flush_list():
        nonlocal list_context
        list_context = None

    while i < len(lines):
        line = lines[i]

        # Em dash strip
        line = line.replace("\u2014", " - ").replace("\u2013", "-")

        # ── Callout block ────────────────────────────────────────────────────
        m = re.match(r'^\[callout:(.+?)\]\s*(.*)', line)
        if m:
            flush_list()
            _add_callout(doc, m.group(1), m.group(2))
            i += 1
            continue

        # ── Fenced code block (``` ... ```) ──────────────────────────────────
        if line.strip().startswith("```"):
            flush_list()
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing ```
            _add_code_block(doc, code_lines)
            continue

        # ── Image block ──────────────────────────────────────────────────────
        m = re.match(r'^\[image:(.+?)\]', line)
        if m:
            flush_list()
            _add_image_block(doc, m.group(1).strip(), images_dir)
            i += 1
            continue

        # ── Horizontal rule ──────────────────────────────────────────────────
        if re.match(r'^\s*---+\s*$', line):
            flush_list()
            _add_hr(doc)
            i += 1
            continue

        # ── Table (collect all rows) ─────────────────────────────────────────
        if _is_table_row(line):
            flush_list()
            table_lines = []
            while i < len(lines) and _is_table_row(lines[i]):
                if not _is_table_separator(lines[i]):
                    table_lines.append(_parse_table_row(lines[i]))
                i += 1
            if table_lines:
                _add_table(doc, table_lines)
            continue

        # ── H2 ───────────────────────────────────────────────────────────────
        m = re.match(r'^##\s+(.*)', line)
        if m:
            flush_list()
            heading_text = m.group(1).strip()
            p = doc.add_paragraph()
            _para_spacing(p, before=160, after=60)
            run = p.add_run(heading_text)
            run.font.size  = Pt(13.5)
            run.font.bold  = True
            run.font.color.rgb = BLUE
            _set_run_font(run)
            i += 1
            continue

        # ── H3 ───────────────────────────────────────────────────────────────
        m = re.match(r'^###\s+(.*)', line)
        if m:
            flush_list()
            heading_text = m.group(1).strip()
            p = doc.add_paragraph()
            _para_spacing(p, before=120, after=40)
            run = p.add_run(heading_text)
            run.font.size  = Pt(11.5)
            run.font.bold  = True
            run.font.color.rgb = BLUE
            _set_run_font(run)
            i += 1
            continue

        # ── Numbered list ────────────────────────────────────────────────────
        m = re.match(r'^(\d+)\.\s+(.*)', line)
        if m:
            list_context = "numbered"
            p = doc.add_paragraph(style="List Number")
            _para_spacing(p, before=20, after=20)
            _add_inline(p, m.group(2).strip(), base_size=10.5)
            i += 1
            continue

        # ── Checkbox list ────────────────────────────────────────────────────
        m = re.match(r'^(\s*)-\s*\[\s*\]\s*(.*)', line)
        if m:
            list_context = "checkbox"
            p = doc.add_paragraph(style="List Bullet")
            _para_spacing(p, before=20, after=20)
            run = p.add_run("\u25a1  ")
            run.font.size = Pt(10.5)
            _set_run_font(run)
            _add_inline(p, m.group(2).strip(), base_size=10.5)
            i += 1
            continue

        # ── Bullet list ──────────────────────────────────────────────────────
        m = re.match(r'^(\s*)-\s+(.*)', line)
        if m:
            list_context = "bullet"
            p = doc.add_paragraph(style="List Bullet")
            _para_spacing(p, before=20, after=20)
            _add_inline(p, m.group(2).strip(), base_size=10.5)
            i += 1
            continue

        # ── Bold label line (e.g. **Layer 1: The Hook**) ─────────────────────
        if line.strip().startswith("**") and line.strip().endswith("**"):
            flush_list()
            inner = line.strip()[2:-2]
            p = doc.add_paragraph()
            _para_spacing(p, before=100, after=30)
            run = p.add_run(inner)
            run.font.size  = Pt(10.5)
            run.font.bold  = True
            run.font.color.rgb = TEXT
            _set_run_font(run)
            i += 1
            continue

        # ── Blank line ───────────────────────────────────────────────────────
        if not line.strip():
            i += 1
            continue

        # ── Regular paragraph ────────────────────────────────────────────────
        flush_list()
        p = doc.add_paragraph()
        _para_spacing(p, before=0, after=60)
        _add_inline(p, line.strip(), base_size=10.5)
        i += 1


# ── Main entry ────────────────────────────────────────────────────────────────

def strip_lead_title(text):
    text = re.sub(r'^#\s+[^\n]+\n', '', text.lstrip(), count=1)
    text = re.sub(r'^\s*##\s+[^\n]+\n', '', text.lstrip(), count=1)
    return text.lstrip()


def strip_raw_cta(text):
    idx = text.lower().find(_CTA_SENTINEL)
    if idx == -1:
        return text
    preceding  = text[:idx]
    last_hr    = preceding.rfind("---")
    if last_hr != -1:
        return text[:last_hr].rstrip()
    return text[:idx].rstrip()


def generate_doc(title, type_, subtitle, content_input, output_path, images_dir=None):
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.2)
        section.right_margin  = Cm(3.2)

    # Default paragraph font
    doc.styles["Normal"].font.name = FONT_NAME
    doc.styles["Normal"].font.size = Pt(10.5)

    # Read content
    raw = Path(content_input).read_text(encoding="utf-8") if os.path.isfile(content_input) else content_input
    raw = strip_lead_title(raw)
    raw = strip_raw_cta(raw)

    # Cover
    _add_cover(doc, title, type_, subtitle)

    # Body
    parse_and_add(doc, raw, images_dir=images_dir)

    # CTA
    _add_cta(doc)

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    if not Path(output_path).exists() or Path(output_path).stat().st_size == 0:
        sys.exit(f"ERROR: DOCX not written to {output_path}")
    print(f"SUCCESS: {output_path}")
    return output_path


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--title",      required=True)
    parser.add_argument("--type",       default="Playbook")
    parser.add_argument("--subtitle",   default="")
    parser.add_argument("--content",    required=True)
    parser.add_argument("--output",     required=True)
    parser.add_argument("--images-dir", default=None)
    args = parser.parse_args()
    generate_doc(
        args.title, args.type, args.subtitle,
        args.content, args.output,
        images_dir=args.images_dir
    )
